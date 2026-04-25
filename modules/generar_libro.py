# -*- coding: utf-8 -*-
from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from ebooklib import epub

# --- IMPORTACIONES PARA PDF ---
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
import html
# ------------------------------

import io
import os
import re
import base64

app = Flask(__name__)

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS COMPARTIDOS
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

def add_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "2E4057")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def _parse_inline(paragraph, text):
    """Convierte Markdown inline a runs de Word: ***negrita+itálica***, **negrita**, *itálica*, ~~tachado~~."""
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*'   # ***bold+italic***
        r'|\*\*(.+?)\*\*'       # **bold**
        r'|\*(.+?)\*'           # *italic*
        r'|__(.+?)__'           # __bold__
        r'|_(.+?)_'             # _italic_
        r'|~~(.+?)~~)'          # ~~strikethrough~~
    )
    ultimo = 0
    for m in pattern.finditer(text):
        if m.start() > ultimo:
            paragraph.add_run(text[ultimo:m.start()])
        full = m.group(0)
        content = next(g for g in m.groups()[1:] if g is not None)
        run = paragraph.add_run(content)
        if full.startswith("***"):
            run.bold = True
            run.italic = True
        elif full.startswith("**") or full.startswith("__"):
            run.bold = True
        elif full.startswith("*") or full.startswith("_"):
            run.italic = True
        elif full.startswith("~~"):
            run.font.strike = True
        ultimo = m.end()
    if ultimo < len(text):
        paragraph.add_run(text[ultimo:])

def _add_page_break(doc):
    """Añade un salto de página explícito."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def _add_heading(doc, text, level):
    """Añade un título usando los estilos semánticos de Word (Heading 1-4) con color corporativo."""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    size_map  = {1: Pt(22), 2: Pt(18), 3: Pt(14), 4: Pt(12)}
    color_map = {
        1: RGBColor(0x1A, 0x25, 0x3A),
        2: RGBColor(0x2E, 0x40, 0x57),
        3: RGBColor(0x2E, 0x40, 0x57),
        4: RGBColor(0x4A, 0x5E, 0x7A),
    }
    p = doc.add_paragraph(style=style_map.get(level, "Heading 1"))
    _parse_inline(p, text)
    for run in p.runs:
        run.font.color.rgb = color_map[level]
        run.font.size = size_map[level]
    return p

def _add_blockquote(doc, text):
    """Añade un bloque de cita con sangría y borde izquierdo azul."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    # Borde izquierdo via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), "2E4057")
    pBdr.append(left)
    pPr.append(pBdr)
    # Texto en itálica y color suave
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x4A, 0x5E, 0x7A)

def _render_table(doc, filas_tabla):
    """Renderiza una tabla Markdown acumulada como tabla Word con estilo."""
    if not filas_tabla:
        return
    num_cols = max(len(f) for f in filas_tabla)
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = "Table Grid"
    for i, fila in enumerate(filas_tabla):
        row = table.add_row()
        for j in range(num_cols):
            celda_texto = fila[j] if j < len(fila) else ""
            cell = row.cells[j]
            cell.text = ""           # limpiar texto default
            add_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run_c = p.add_run(celda_texto)
            if i == 0:               # fila de cabecera
                run_c.bold = True
                run_c.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_background(cell, "2E4057")
            else:
                run_c.font.color.rgb = RGBColor(0x1A, 0x25, 0x3A)
                if i % 2 == 0:
                    set_cell_background(cell, "EEF2F7")
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR DOCX — FUNCIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

def generar_docx(titulo, contenido, portada_b64=""):
    doc = Document()

    # ── 1. Página de portada: imagen a página completa, sin márgenes ─────────
    section_portada = doc.sections[0]
    section_portada.page_width      = Cm(21)
    section_portada.page_height     = Cm(29.7)
    section_portada.top_margin      = Cm(0)
    section_portada.bottom_margin   = Cm(0)
    section_portada.left_margin     = Cm(0)
    section_portada.right_margin    = Cm(0)

    if portada_b64:
        try:
            if "," in portada_b64:
                portada_b64 = portada_b64.split(",", 1)[1]
            img_data   = base64.b64decode(portada_b64)
            img_stream = io.BytesIO(img_data)
            p_cover = doc.add_paragraph()
            p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cover.paragraph_format.space_before = Pt(0)
            p_cover.paragraph_format.space_after  = Pt(0)
            run_cover = p_cover.add_run()
            # Imagen que cubre la página completa (A4: 21 × 29.7 cm)
            run_cover.add_picture(img_stream, width=Cm(21), height=Cm(29.7))
        except Exception as e:
            print(f"[DOCX] Error procesando portada: {e}")

    # ── 2. Nueva sección con márgenes normales de libro ───────────────────────
    section_content = doc.add_section(WD_SECTION.NEW_PAGE)
    section_content.page_width    = Cm(21)
    section_content.page_height   = Cm(29.7)
    section_content.top_margin    = Cm(2.5)
    section_content.bottom_margin = Cm(2.5)
    section_content.left_margin   = Cm(3.0)
    section_content.right_margin  = Cm(2.5)

    # ── 3. Página de título ───────────────────────────────────────────────────
    # Espacio vertical centrado
    for _ in range(8):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(0)

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(28)
    run_titulo.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # Línea decorativa bajo el título
    p_linea = doc.add_paragraph()
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_linea = p_linea.add_run("─" * 30)
    run_linea.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    run_linea.font.size = Pt(12)

    _add_page_break(doc)

    # ── 4. Procesar secciones del libro (separadas por \n---\n) ──────────────
    secciones = re.split(r"\n---\n", contenido)

    for idx, seccion in enumerate(secciones):
        lineas      = seccion.strip().split("\n")
        en_tabla    = False
        filas_tabla = []

        for linea in lineas:
            linea = linea.strip()
            if not linea:
                # Línea vacía: cierra tabla pendiente si la hay
                if en_tabla and filas_tabla:
                    _render_table(doc, filas_tabla)
                    filas_tabla = []
                    en_tabla    = False
                continue

            # ── Detección de tabla Markdown ───────────────────────────────
            if "|" in linea:
                if re.match(r"^\|[-| :]+\|$", linea):   # separador de cabecera
                    continue
                celdas = [c.strip() for c in linea.strip("|").split("|")]
                filas_tabla.append(celdas)
                en_tabla = True
                continue
            else:
                if en_tabla and filas_tabla:
                    _render_table(doc, filas_tabla)
                    filas_tabla = []
                    en_tabla    = False

            # ── Títulos / Headings ────────────────────────────────────────
            if linea.startswith("#### "):
                _add_heading(doc, linea[5:], 4)
            elif linea.startswith("### "):
                _add_heading(doc, linea[4:], 3)
            elif linea.startswith("## "):
                _add_heading(doc, linea[3:], 2)
            elif linea.startswith("# "):
                _add_heading(doc, linea[2:], 1)

            # ── Cita / Blockquote ─────────────────────────────────────────
            elif linea.startswith("> "):
                _add_blockquote(doc, linea[2:])

            # ── Listas ────────────────────────────────────────────────────
            elif linea.startswith("- ") or linea.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _parse_inline(p, linea[2:])

            elif re.match(r"^\d+\.\s", linea):
                p = doc.add_paragraph(style="List Number")
                _parse_inline(p, re.sub(r"^\d+\.\s", "", linea))

            # ── Párrafo normal ────────────────────────────────────────────
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)
                _parse_inline(p, linea)

        # Cerrar tabla pendiente al final de la sección
        if en_tabla and filas_tabla:
            _render_table(doc, filas_tabla)

        # Salto de página entre capítulos (no después del último)
        if idx < len(secciones) - 1:
            _add_page_break(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ══════════════════════════════════════════════════════════════════════════════
# RUTAS FLASK
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})

@app.route("/generar/docx", methods=["POST"])
def generar():
    data        = request.get_json(force=True)
    titulo      = data.get("titulo", "Libro")
    contenido   = data.get("contenido", "")
    portada_b64 = data.get("portada_base64", "")   # ← nuevo campo
    buffer = generar_docx(titulo, contenido, portada_b64)
    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{titulo}.docx",
    )


# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR EPUB (sin cambios)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/generar/epub", methods=["POST"])
def generar_epub():
    data        = request.json
    titulo      = data.get("titulo", "Libro_Generado")
    contenido   = data.get("contenido", "")
    portada_b64 = data.get("portada_base64", "")

    book = epub.EpubBook()
    book.set_identifier("id123456")
    book.set_title(titulo)
    book.set_language("es")

    if portada_b64:
        if "," in portada_b64:
            portada_b64 = portada_b64.split(",", 1)[1]
        cover_data = base64.b64decode(portada_b64)
        book.set_cover("portada.png", cover_data)

    epub_chapters   = []
    texto_completo  = contenido.replace("\n\n---\n\n", "\n")
    lines           = texto_completo.split("\n")

    current_title   = "Inicio"
    current_html    = ""
    chapter_counter = 0
    in_table        = False
    table_html      = ""

    for line in lines:
        line = line.strip()

        if not line:
            if in_table:
                in_table    = False
                current_html += table_html + "</table><br/>"
                table_html  = ""
            continue

        if line.startswith("|"):
            if not in_table:
                in_table   = True
                table_html = "<table style='border-collapse:collapse;width:100%;margin-bottom:1em;border:1px solid #ccc;'>"
            if re.match(r"^[\s\|\-\:]+$", line):
                continue
            celdas = [c.strip() for c in line.strip("|").split("|")]
            table_html += "<tr>"
            for celda in celdas:
                if "<tr>" not in table_html:
                    table_html += f"<th style='padding:8px;background-color:#f2f2f2;border:1px solid #ccc;text-align:left;'>{celda}</th>"
                else:
                    table_html += f"<td style='padding:8px;border:1px solid #ccc;'>{celda}</td>"
            table_html += "</tr>"
            continue

        if in_table:
            in_table     = False
            current_html += table_html + "</table><br/>"
            table_html   = ""

        is_h3_markdown = line.startswith("###")
        clean_line     = re.sub(r"^#+\s*", "", line).replace("**", "").replace("*", "").strip()
        is_main_title  = re.match(
            r"^(introducción|conclusión|recursos prácticos|apéndice|capítulo\s+\d+)",
            clean_line, re.I
        )
        is_sub_title = is_h3_markdown or re.match(
            r"^[A-ZÁÉÍÓÚÑ0-9][^.]{5,60}:$|^[A-ZÁÉÍÓÚÑ0-9][^.]{3,40}$", clean_line
        )

        if is_main_title:
            if current_html:
                c = epub.EpubHtml(title=current_title, file_name=f"cap_{chapter_counter}.xhtml", lang="es")
                c.content = f"<html><head></head><body>{current_html}</body></html>"
                epub_chapters.append(c)
                book.add_item(c)
                chapter_counter += 1
            current_title = clean_line
            current_html  = f"<h1>{clean_line}</h1>"
        elif is_sub_title:
            current_html += f"<h2>{clean_line}</h2>"
        elif line.startswith("> "):
            # NUEVO: Manejo de blockquotes para EPUB
            clean_quote = clean_line[1:].strip() if clean_line.startswith(">") else clean_line
            current_html += f"<blockquote style='margin: 1.5em 2em; padding-left: 1em; border-left: 4px solid #2E4057; font-style: italic; color: #4A5E7A;'>{clean_quote}</blockquote>"
        else:
            current_html += f"<p>{clean_line}</p>"

    if in_table:
        current_html += table_html + "</table>"
    if current_html:
        c = epub.EpubHtml(title=current_title, file_name=f"cap_{chapter_counter}.xhtml", lang="es")
        c.content = f"<html><head></head><body>{current_html}</body></html>"
        epub_chapters.append(c)
        book.add_item(c)

    book.toc   = tuple(epub_chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav"] + epub_chapters

    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{titulo}.epub",
        mimetype="application/epub+zip",
    )


# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR PDF (sin cambios)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/generar/pdf", methods=["POST"])
def generar_pdf_route():
    data        = request.json
    titulo      = data.get("titulo", "Libro_Generado")
    contenido   = data.get("contenido", "")
    portada_b64 = data.get("portada_base64", "")

    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    Story  = []
    styles = getSampleStyleSheet()

    title_style  = ParagraphStyle(name="TitleStyle",  parent=styles["Heading1"], alignment=TA_CENTER, fontSize=24, spaceAfter=20)
    h1_style     = ParagraphStyle(name="H1Style",     parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h2_style     = ParagraphStyle(name="H2Style",     parent=styles["Heading2"], fontSize=14, spaceAfter=10)
    normal_style = ParagraphStyle(name="NormalStyle", parent=styles["Normal"],   alignment=TA_JUSTIFY, spaceAfter=8, fontSize=11, leading=14)
    
    # NUEVO: Estilo para las citas textuales
    blockquote_style = ParagraphStyle(
        name="BlockquoteStyle",
        parent=styles["Normal"],
        alignment=TA_JUSTIFY,
        leftIndent=30,
        rightIndent=30,
        fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#4A5E7A"),
        spaceBefore=12,
        spaceAfter=12,
        leading=14
    )

    cover_page_func = None

    if portada_b64:
        try:
            if "," in portada_b64:
                portada_b64 = portada_b64.split(",", 1)[1]
            img_data   = base64.b64decode(portada_b64)
            img_stream = io.BytesIO(img_data)
            cover_img  = ImageReader(img_stream)

            # NUEVO: Función callback para dibujar la imagen a pantalla completa (Full Bleed)
            def cover_page(canvas, doc):
                canvas.saveState()
                canvas.drawImage(cover_img, 0, 0, width=A4[0], height=A4[1], preserveAspectRatio=False)
                canvas.restoreState()

            cover_page_func = cover_page

            # Agregamos una página en blanco al Story. 
            # Esto empuja el texto al inicio de la página 2, dejando la portada sola.
            Story.append(Spacer(1, 1))
            Story.append(PageBreak())
        except Exception as e:
            print(f"Error procesando la portada: {e}")

    Story.append(Spacer(1, 2 * inch))
    Story.append(Paragraph(titulo, title_style))
    Story.append(PageBreak())

    texto_completo = contenido.replace("\n\n---\n\n", "\n")
    lines          = texto_completo.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            Story.append(Spacer(1, 0.1 * inch))
            continue

        if line.startswith("* ") or line.startswith("- "):
            line = "• " + line[2:]

        escaped_line   = html.escape(line)
        formatted_line = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", escaped_line)
        formatted_line = re.sub(r"\*\*(.+?)\*\*",     r"<b>\1</b>",         formatted_line)
        formatted_line = re.sub(r"(?<!\*)\*(?!\s)([^\*]+?)(?<!\s)\*(?!\*)", r"<i>\1</i>", formatted_line)
        formatted_line = re.sub(r"_(?!\s)([^_]+?)(?<!\s)_",                 r"<i>\1</i>", formatted_line)

        clean_line_no_stars = re.sub(r"^#+\s*", "", line).replace("**", "").replace("*", "").strip()
        is_main_title = re.match(
            r"^(introducción|conclusión|recursos prácticos|apéndice|capítulo\s+\d+)",
            clean_line_no_stars, re.I
        )
        formatted_line = re.sub(r"^#+\s*", "", formatted_line)

        if is_main_title:
            Story.append(PageBreak())
            Story.append(Paragraph(formatted_line, h1_style))
        elif line.startswith("###") or line.startswith("##"):
            Story.append(Paragraph(formatted_line, h2_style))
        elif line.startswith("> "):
            # NUEVO: Como html.escape convierte '>' a '&gt;', lo limpiamos aquí
            clean_quote_pdf = formatted_line.replace("&gt; ", "", 1).strip()
            Story.append(Paragraph(clean_quote_pdf, blockquote_style))
        else:
            Story.append(Paragraph(formatted_line, normal_style))

    # NUEVO: Compilamos usando la función onFirstPage si existe una portada
    if cover_page_func:
        doc.build(Story, onFirstPage=cover_page_func)
    else:
        doc.build(Story)
        
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{titulo}.pdf",
        mimetype="application/pdf",
    )

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR DESCRIPCIÓN HTML (.txt)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/generar/descripcion", methods=["POST"])
def generar_descripcion():
    data = request.json
    titulo = data.get("titulo", "Libro")
    descripcion = data.get("descripcion", "")

    # Formatear el texto plano a un HTML estructurado y limpio
    html_content = f"<h2>{titulo}</h2>\n"
    
    # Separar por dobles saltos de línea para crear párrafos HTML
    parrafos = descripcion.split('\n\n')
    for p in parrafos:
        p_limpio = p.strip()
        if p_limpio:
            html_content += f"<p>{p_limpio}</p>\n"

    buffer = io.BytesIO()
    buffer.write(html_content.encode('utf-8'))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{titulo}_descripcion.txt",
        mimetype="text/plain"
    )

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR MUESTRA GRATIS PDF (Portada + Intro)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/generar/muestra_pdf", methods=["POST"])
def generar_muestra_pdf():
    data        = request.json
    titulo      = data.get("titulo", "Muestra_Generada")
    intro       = data.get("intro", "")
    portada_b64 = data.get("portada_base64", "")

    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    Story  = []
    styles = getSampleStyleSheet()

    # Estilos
    title_style  = ParagraphStyle(name="TitleStyle",  parent=styles["Heading1"], alignment=TA_CENTER, fontSize=24, spaceAfter=20)
    h1_style     = ParagraphStyle(name="H1Style",     parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h2_style     = ParagraphStyle(name="H2Style",     parent=styles["Heading2"], fontSize=14, spaceAfter=10)
    normal_style = ParagraphStyle(name="NormalStyle", parent=styles["Normal"],   alignment=TA_JUSTIFY, spaceAfter=8, fontSize=11, leading=14)
    blockquote_style = ParagraphStyle(
        name="BlockquoteStyle", parent=styles["Normal"], alignment=TA_JUSTIFY,
        leftIndent=30, rightIndent=30, fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#4A5E7A"), spaceBefore=12, spaceAfter=12, leading=14
    )

    cover_page_func = None

    # Procesar Portada a pantalla completa
    if portada_b64:
        try:
            if "base64," in portada_b64:
                portada_b64 = portada_b64.split("base64,")[1]
            img_data   = base64.b64decode(portada_b64)
            img_stream = io.BytesIO(img_data)
            cover_img  = ImageReader(img_stream)

            def cover_page(canvas, doc):
                canvas.saveState()
                canvas.drawImage(cover_img, 0, 0, width=A4[0], height=A4[1], preserveAspectRatio=False)
                canvas.restoreState()

            cover_page_func = cover_page
            Story.append(Spacer(1, 1))
            Story.append(PageBreak())
        except Exception as e:
            print(f"Error procesando la portada de la muestra: {e}")

    # Título de la muestra
    Story.append(Spacer(1, 2 * inch))
    Story.append(Paragraph(f"Muestra Gratuita: {titulo}", title_style))
    Story.append(PageBreak())

    # Procesar el contenido de la introducción
    texto_completo = intro.replace("\n\n---\n\n", "\n")
    lines          = texto_completo.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            Story.append(Spacer(1, 0.1 * inch))
            continue

        # Convertir Markdown a HTML compatible con ReportLab
        escaped_line   = html.escape(line)
        formatted_line = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", escaped_line)
        formatted_line = re.sub(r"\*\*(.+?)\*\*",     r"<b>\1</b>",         formatted_line)
        formatted_line = re.sub(r"(?<!\*)\*(?!\s)([^\*]+?)(?<!\s)\*(?!\*)", r"<i>\1</i>", formatted_line)
        formatted_line = re.sub(r"_(?!\s)([^_]+?)(?<!\s)_",                 r"<i>\1</i>", formatted_line)

        clean_line_no_stars = re.sub(r"^#+\s*", "", line).replace("**", "").replace("*", "").strip()
        is_main_title = re.match(
            r"^(introducción|conclusión|recursos prácticos|apéndice|capítulo\s+\d+)",
            clean_line_no_stars, re.I
        )
        formatted_line = re.sub(r"^#+\s*", "", formatted_line)

        if is_main_title:
            Story.append(PageBreak())
            Story.append(Paragraph(formatted_line, h1_style))
        elif line.startswith("###") or line.startswith("##"):
            Story.append(Paragraph(formatted_line, h2_style))
        elif line.startswith("> "):
            clean_quote_pdf = formatted_line.replace("&gt; ", "", 1).strip()
            Story.append(Paragraph(clean_quote_pdf, blockquote_style))
        else:
            Story.append(Paragraph(formatted_line, normal_style))

    if cover_page_func:
        doc.build(Story, onFirstPage=cover_page_func)
    else:
        doc.build(Story)
        
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"Muestra_{titulo}.pdf",
        mimetype="application/pdf",
    )

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR DE RESUMEN PDF
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/generar/resumen_pdf", methods=["POST"])
def generar_resumen_pdf():
    data        = request.json
    titulo      = data.get("titulo", "Resumen_Generado")
    resumen     = data.get("resumen", "")
    portada_b64 = data.get("portada_base64", "")

    buffer = io.BytesIO()
    doc    = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    Story  = []
    styles = getSampleStyleSheet()

    # Estilos heredados del libro principal
    title_style  = ParagraphStyle(name="TitleStyle",  parent=styles["Heading1"], alignment=TA_CENTER, fontSize=24, spaceAfter=20)
    h1_style     = ParagraphStyle(name="H1Style",     parent=styles["Heading1"], fontSize=18, spaceAfter=12)
    h2_style     = ParagraphStyle(name="H2Style",     parent=styles["Heading2"], fontSize=14, spaceAfter=10)
    normal_style = ParagraphStyle(name="NormalStyle", parent=styles["Normal"],   alignment=TA_JUSTIFY, spaceAfter=8, fontSize=11, leading=14)
    blockquote_style = ParagraphStyle(
        name="BlockquoteStyle", parent=styles["Normal"], alignment=TA_JUSTIFY,
        leftIndent=30, rightIndent=30, fontName="Helvetica-Oblique",
        textColor=colors.HexColor("#4A5E7A"), spaceBefore=12, spaceAfter=12, leading=14
    )

    cover_page_func = None

    # Procesar Portada a pantalla completa (Full Bleed)
    if portada_b64:
        try:
            # Limpieza del prefijo de Base64 si existe
            if "base64," in portada_b64:
                portada_b64 = portada_b64.split("base64,")[1]
            elif "," in portada_b64:
                portada_b64 = portada_b64.split(",", 1)[1]
                
            img_data   = base64.b64decode(portada_b64)
            img_stream = io.BytesIO(img_data)
            cover_img  = ImageReader(img_stream)

            def cover_page(canvas, doc):
                canvas.saveState()
                canvas.drawImage(cover_img, 0, 0, width=A4[0], height=A4[1], preserveAspectRatio=False)
                canvas.restoreState()

            cover_page_func = cover_page
            Story.append(Spacer(1, 1))
            Story.append(PageBreak())
        except Exception as e:
            print(f"Error procesando la portada del resumen: {e}")

    # Título del resumen
    Story.append(Spacer(1, 2 * inch))
    Story.append(Paragraph(titulo, title_style))
    Story.append(PageBreak())

    # Procesar el contenido del resumen
    texto_completo = resumen.replace("\n\n---\n\n", "\n")
    lines          = texto_completo.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            Story.append(Spacer(1, 0.1 * inch))
            continue

        # Formatear viñetas
        if line.startswith("* ") or line.startswith("- "):
            line = "• " + line[2:]

        escaped_line   = html.escape(line)
        
        # Parseo robusto de Markdown (previniendo el Parse error de ReportLab)
        formatted_line = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", escaped_line)
        formatted_line = re.sub(r"\*\*(.+?)\*\*",     r"<b>\1</b>",         formatted_line)
        formatted_line = re.sub(r"(?<!\*)\*(?!\s)([^\*]+?)(?<!\s)\*(?!\*)", r"<i>\1</i>", formatted_line)
        formatted_line = re.sub(r"_(?!\s)([^_]+?)(?<!\s)_",                 r"<i>\1</i>", formatted_line)

        # Identificamos si es un título principal para el salto de página
        clean_line_no_stars = re.sub(r"^#+\s*", "", line).replace("**", "").replace("*", "").strip()
        
        # Ampliamos la regex para capturar "índice" u otras palabras clave comunes en resúmenes
        is_main_title = re.match(
            r"^(introducción|conclusión|recursos prácticos|apéndice|capítulo\s+\d+|índice)",
            clean_line_no_stars, re.I
        )
        formatted_line = re.sub(r"^#+\s*", "", formatted_line)

        if is_main_title:
            Story.append(PageBreak())
            Story.append(Paragraph(formatted_line, h1_style))
        elif line.startswith("###") or line.startswith("##"):
            Story.append(Paragraph(formatted_line, h2_style))
        elif line.startswith("&gt; ") or line.startswith("> "):
            # Limpiamos tanto la forma normal como la escapada en HTML del Blockquote
            clean_quote_pdf = formatted_line.replace("&gt; ", "", 1).replace("> ", "", 1).strip()
            Story.append(Paragraph(clean_quote_pdf, blockquote_style))
        else:
            Story.append(Paragraph(formatted_line, normal_style))

    # Compilar aplicando la portada si existe
    if cover_page_func:
        doc.build(Story, onFirstPage=cover_page_func)
    else:
        doc.build(Story)
        
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{titulo}.pdf",
        mimetype="application/pdf",
    )
    
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 9000))
    print(f"Servidor iniciado en http://127.0.0.1:{port}")
    app.run(host="0.0.0.0", port=port, debug=False)
